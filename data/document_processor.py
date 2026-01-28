"""
Document Processor for Financial Research Agent.

Handles file upload processing and text extraction for various document types.
Supports: PDF, Word (docx), Text, Markdown, CSV, Excel, HTML
"""

import io
import re
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass

logger = logging.getLogger(__name__)

# Supported file types and their MIME types
SUPPORTED_TYPES = {
    "pdf": ["application/pdf"],
    "docx": ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"],
    "doc": ["application/msword"],
    "txt": ["text/plain"],
    "md": ["text/markdown", "text/x-markdown"],
    "csv": ["text/csv"],
    "xlsx": ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"],
    "xls": ["application/vnd.ms-excel"],
    "html": ["text/html"],
    "htm": ["text/html"],
}

# File extensions to type mapping
EXTENSION_TO_TYPE = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "doc",
    ".txt": "txt",
    ".md": "md",
    ".markdown": "md",
    ".csv": "csv",
    ".xlsx": "xlsx",
    ".xls": "xls",
    ".html": "html",
    ".htm": "html",
}


@dataclass
class ProcessedDocument:
    """Result of processing a document."""
    text: str
    metadata: Dict[str, Any]
    file_type: str
    page_count: Optional[int] = None
    word_count: int = 0
    char_count: int = 0
    
    def __post_init__(self):
        if self.text:
            self.word_count = len(self.text.split())
            self.char_count = len(self.text)


class DocumentProcessor:
    """
    Processes various document types and extracts text content.
    
    Supports:
    - PDF: Uses PyPDF2 or pdfplumber
    - Word (docx): Uses python-docx
    - Text/Markdown: Direct read
    - CSV/Excel: pandas conversion to text
    - HTML: BeautifulSoup extraction
    """
    
    def __init__(self):
        """Initialize the document processor."""
        self._check_dependencies()
    
    def _check_dependencies(self) -> Dict[str, bool]:
        """Check which optional dependencies are available."""
        deps = {}
        
        try:
            import PyPDF2
            deps["pypdf2"] = True
        except ImportError:
            deps["pypdf2"] = False
        
        try:
            import pdfplumber
            deps["pdfplumber"] = True
        except ImportError:
            deps["pdfplumber"] = False
        
        try:
            import docx
            deps["python-docx"] = True
        except ImportError:
            deps["python-docx"] = False
        
        try:
            from bs4 import BeautifulSoup
            deps["beautifulsoup"] = True
        except ImportError:
            deps["beautifulsoup"] = False
        
        try:
            import pandas
            deps["pandas"] = True
        except ImportError:
            deps["pandas"] = False
        
        self._dependencies = deps
        return deps
    
    def get_file_type(self, filename: str, mime_type: Optional[str] = None) -> Optional[str]:
        """
        Determine file type from filename or MIME type.
        
        Args:
            filename: Original filename
            mime_type: Optional MIME type
            
        Returns:
            File type string or None if unsupported
        """
        # Try extension first
        ext = Path(filename).suffix.lower()
        if ext in EXTENSION_TO_TYPE:
            return EXTENSION_TO_TYPE[ext]
        
        # Try MIME type
        if mime_type:
            for file_type, mime_types in SUPPORTED_TYPES.items():
                if mime_type in mime_types:
                    return file_type
        
        return None
    
    def is_supported(self, filename: str, mime_type: Optional[str] = None) -> bool:
        """Check if a file type is supported."""
        return self.get_file_type(filename, mime_type) is not None
    
    def process(
        self,
        content: bytes,
        filename: str,
        mime_type: Optional[str] = None
    ) -> ProcessedDocument:
        """
        Process a document and extract text.
        
        Args:
            content: File content as bytes
            filename: Original filename
            mime_type: Optional MIME type
            
        Returns:
            ProcessedDocument with extracted text and metadata
        """
        file_type = self.get_file_type(filename, mime_type)
        
        if not file_type:
            raise ValueError(f"Unsupported file type: {filename}")
        
        metadata = {
            "filename": filename,
            "file_type": file_type,
            "size_bytes": len(content),
        }
        
        try:
            if file_type == "pdf":
                text, page_count = self._process_pdf(content)
                metadata["page_count"] = page_count
                return ProcessedDocument(text=text, metadata=metadata, file_type=file_type, page_count=page_count)
            
            elif file_type in ("docx", "doc"):
                text = self._process_docx(content)
                return ProcessedDocument(text=text, metadata=metadata, file_type=file_type)
            
            elif file_type in ("txt", "md"):
                text = self._process_text(content)
                return ProcessedDocument(text=text, metadata=metadata, file_type=file_type)
            
            elif file_type == "csv":
                text = self._process_csv(content)
                return ProcessedDocument(text=text, metadata=metadata, file_type=file_type)
            
            elif file_type in ("xlsx", "xls"):
                text = self._process_excel(content, file_type)
                return ProcessedDocument(text=text, metadata=metadata, file_type=file_type)
            
            elif file_type in ("html", "htm"):
                text = self._process_html(content)
                return ProcessedDocument(text=text, metadata=metadata, file_type=file_type)
            
            else:
                raise ValueError(f"No processor for file type: {file_type}")
                
        except Exception as e:
            logger.error(f"Error processing {filename}: {e}")
            raise
    
    def _process_pdf(self, content: bytes) -> Tuple[str, int]:
        """Extract text from PDF."""
        # Try pdfplumber first (better quality)
        if self._dependencies.get("pdfplumber"):
            try:
                import pdfplumber
                
                text_parts = []
                page_count = 0
                
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    page_count = len(pdf.pages)
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                
                return "\n\n".join(text_parts), page_count
            except Exception as e:
                logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
        
        # Fallback to PyPDF2
        if self._dependencies.get("pypdf2"):
            try:
                import PyPDF2
                
                text_parts = []
                
                reader = PyPDF2.PdfReader(io.BytesIO(content))
                page_count = len(reader.pages)
                
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                return "\n\n".join(text_parts), page_count
            except Exception as e:
                logger.error(f"PyPDF2 failed: {e}")
                raise
        
        raise ImportError("No PDF library available. Install pdfplumber or PyPDF2.")
    
    def _process_docx(self, content: bytes) -> str:
        """Extract text from Word document."""
        if not self._dependencies.get("python-docx"):
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        import docx
        
        doc = docx.Document(io.BytesIO(content))
        
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts)
    
    def _process_text(self, content: bytes) -> str:
        """Process plain text or markdown file."""
        # Try UTF-8 first, then fall back to other encodings
        for encoding in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        
        # Last resort: decode with replacement
        return content.decode("utf-8", errors="replace")
    
    def _process_csv(self, content: bytes) -> str:
        """Convert CSV to text representation."""
        if not self._dependencies.get("pandas"):
            # Fallback: simple text representation
            text = self._process_text(content)
            return f"CSV Data:\n{text}"
        
        import pandas as pd
        
        try:
            df = pd.read_csv(io.BytesIO(content))
            
            # Generate text representation
            parts = []
            parts.append(f"CSV with {len(df)} rows and {len(df.columns)} columns")
            parts.append(f"Columns: {', '.join(df.columns.tolist())}")
            parts.append("")
            
            # Add summary statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                parts.append("Summary Statistics:")
                parts.append(df[numeric_cols].describe().to_string())
                parts.append("")
            
            # Add first few rows as sample
            parts.append("Sample Data (first 10 rows):")
            parts.append(df.head(10).to_string())
            
            return "\n".join(parts)
        except Exception as e:
            logger.warning(f"Pandas CSV processing failed: {e}")
            return self._process_text(content)
    
    def _process_excel(self, content: bytes, file_type: str) -> str:
        """Convert Excel to text representation."""
        if not self._dependencies.get("pandas"):
            raise ImportError("pandas not installed. Run: pip install pandas openpyxl")
        
        import pandas as pd
        
        try:
            # Read all sheets
            engine = "openpyxl" if file_type == "xlsx" else "xlrd"
            
            try:
                sheets = pd.read_excel(io.BytesIO(content), sheet_name=None, engine=engine)
            except Exception:
                # Try without specifying engine
                sheets = pd.read_excel(io.BytesIO(content), sheet_name=None)
            
            parts = []
            parts.append(f"Excel file with {len(sheets)} sheet(s)")
            parts.append("")
            
            for sheet_name, df in sheets.items():
                parts.append(f"=== Sheet: {sheet_name} ===")
                parts.append(f"{len(df)} rows, {len(df.columns)} columns")
                parts.append(f"Columns: {', '.join(str(c) for c in df.columns.tolist())}")
                
                # Summary for numeric columns
                numeric_cols = df.select_dtypes(include=['number']).columns
                if len(numeric_cols) > 0:
                    parts.append("\nSummary Statistics:")
                    parts.append(df[numeric_cols].describe().to_string())
                
                # Sample data
                parts.append("\nSample Data:")
                parts.append(df.head(10).to_string())
                parts.append("")
            
            return "\n".join(parts)
        except Exception as e:
            logger.error(f"Excel processing failed: {e}")
            raise
    
    def _process_html(self, content: bytes) -> str:
        """Extract text from HTML."""
        html_text = self._process_text(content)
        
        if self._dependencies.get("beautifulsoup"):
            try:
                from bs4 import BeautifulSoup
                
                soup = BeautifulSoup(html_text, "html.parser")
                
                # Remove script and style elements
                for element in soup(["script", "style", "nav", "footer", "header"]):
                    element.decompose()
                
                # Get text
                text = soup.get_text(separator="\n")
                
                # Clean up whitespace
                lines = [line.strip() for line in text.splitlines()]
                text = "\n".join(line for line in lines if line)
                
                return text
            except Exception as e:
                logger.warning(f"BeautifulSoup processing failed: {e}")
        
        # Fallback: simple regex-based extraction
        # Remove tags
        text = re.sub(r'<[^>]+>', ' ', html_text)
        # Clean up whitespace
        text = re.sub(r'\s+', ' ', text)
        return text.strip()
    
    def chunk_text(
        self,
        text: str,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Dict[str, Any]]:
        """
        Chunk text for RAG indexing.
        
        Args:
            text: Text to chunk
            chunk_size: Target chunk size in characters
            chunk_overlap: Overlap between chunks
            metadata: Metadata to include with each chunk
            
        Returns:
            List of chunk dictionaries with 'content' and 'metadata'
        """
        if not text:
            return []
        
        # Clean text
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        
        chunks = []
        start = 0
        chunk_index = 0
        
        while start < len(text):
            # Find end of chunk
            end = start + chunk_size
            
            if end < len(text):
                # Try to break at paragraph
                para_break = text.rfind('\n\n', start, end)
                if para_break > start + chunk_size // 2:
                    end = para_break + 2
                else:
                    # Try to break at sentence
                    sentence_break = text.rfind('. ', start, end)
                    if sentence_break > start + chunk_size // 2:
                        end = sentence_break + 2
            
            chunk_text = text[start:end].strip()
            
            if chunk_text:
                chunk_metadata = {
                    "chunk_index": chunk_index,
                    **(metadata or {})
                }
                
                chunks.append({
                    "content": chunk_text,
                    "metadata": chunk_metadata
                })
                chunk_index += 1
            
            # Move start with overlap
            start = end - chunk_overlap
            if start <= 0:
                start = end
        
        return chunks


# Global instance
_document_processor: Optional[DocumentProcessor] = None


def get_document_processor() -> DocumentProcessor:
    """Get or create the global document processor instance."""
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessor()
    return _document_processor
